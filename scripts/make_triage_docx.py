# -*- coding: utf-8 -*-
"""Masaüstüne triage anlatım Word dokümanı üretir."""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

out = Path(r"C:\Users\baran\Desktop") / "T2DM_Chatbot_Triage_Katmani_Detayli_Anlatim.docx"
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def set_run_font(run, size=11, bold=False, code=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Consolas" if code else "Calibri"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    name = "Consolas" if code else "Calibri"
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_p(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 11, bold=bold)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, 11)
    return p


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    set_run_font(run, 9, code=True, color=(40, 40, 40))
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_run_font(cell.paragraphs[0].add_run(h), 10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            set_run_font(cell.paragraphs[0].add_run(str(val)), 9)
    doc.add_paragraph()


# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(title.add_run("Tip-2 Diyabet Chatbot — Triage Katmanı"), 22, bold=True)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    sub.add_run("Teorik tasarım, mimari, dizin yapısı ve implementasyon anlatımı"),
    13,
)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    meta.add_run(
        "Proje: Type 2 Diabet Chatbot  |  4 Ağustos 2026  |  "
        "Katmanlar: Adım 1–3 (Numeric + Regex + Fusion/FT/Grey-zone)"
    ),
    10,
    color=(80, 80, 80),
)

add_p(
    "Bu belge, aynı gün içinde tamamlanan hibrit güvenlik / aciliyet (triage) "
    "katmanını sonradan hatırlanabilir şekilde anlatır: neden böyle tasarlandı, "
    "klasörler nerede, her .py ne işe yarar, kritik kod parçaları ne demek."
)

doc.add_heading("1. Proje bağlamı — chatbot ne yapıyor?", level=1)
add_p(
    "Ürün: Türkçe Tip-2 diyabet hasta eğitim asistanı. RAG ile doğrulanmış "
    "kaynaklardan cevap üretir (bge-m3 retrieve → mmarco rerank → NVIDIA Nemotron "
    "chat). Triage katmanı RAG’tan ÖNCE çalışır: acil / doz-tanı reddi / sarı uyarı "
    "/ yeşil eğitim ayrımını yapar. Amaç under-triage’ı (acil kaçırma) pahalı, "
    "over-triage’ı (gereksiz 112) daha kabul edilebilir tutmaktır."
)
add_p("Dört seviye (API sözleşmesi):", bold=True)
add_bullet("EMERGENCY — 112 / acil servis canned mesajı (RAG yok)")
add_bullet("REFUSE — doz / tanı / reçete / jailbreak reddi (112 yok, ayrı canned)")
add_bullet("YELLOW — RAG devam + hekim uyarısı (veya tempered grey-zone metni)")
add_bullet("GREEN — normal RAG + LLM eğitimi")

doc.add_heading("2. Proje dizini — triage’a odak", level=1)
add_p("Kök: Desktop/Staj/Type 2 Diabet Chatbot. Triage ile ilgili iskelet:")
add_code(
    "Type 2 Diabet Chatbot/\n"
    "├── src/api/\n"
    "│   ├── app.py              # FastAPI /chat\n"
    "│   ├── pipeline.py         # triage → retrieve → rerank → LLM\n"
    "│   ├── llm.py              # Nemotron chat (hasta cevabı)\n"
    "│   └── triage/             # ★ HİBRİT TRİAGE PAKETİ\n"
    "│       ├── __init__.py     # detect_triage / detailed / canned\n"
    "│       ├── text_utils.py   # ortak norm()\n"
    "│       ├── numeric.py      # Adım 1 — glukoz sayıları\n"
    "│       ├── regex_flags.py  # Adım 2 — hard + soft regex\n"
    "│       ├── ft_encoder.py   # Adım 3 — bge-m3 örtük skor\n"
    "│       ├── fusion.py       # Adım 3 — skor + band + guard\n"
    "│       └── grey_zone.py    # Adım 3 — Nemotron JSON sınıflandırma\n"
    "├── src/retrieval/embed.py  # bge-m3 (FT encoder reuse)\n"
    "├── src/eval/\n"
    "│   ├── check_numeric_triage.py\n"
    "│   ├── check_regex_triage.py\n"
    "│   ├── check_fusion.py\n"
    "│   └── tune_band.py\n"
    "├── data/gold/gold_set.jsonl\n"
    "└── frontend/               # Vite UI (triage_level rozeti)"
)

doc.add_heading("2.1 Dosya-dosya ne işe yarar?", level=2)
add_table(
    ["Dosya", "Görevi"],
    [
        ["text_utils.py", "Türkçe normalizasyon (İ güvenli). Numeric indeks + regex ortak kaynak."],
        ["numeric.py", "Sayı çıkar, glukoz bağlamı, DKA/bilinç/dehidratasyon, eşik ağacı."],
        ["regex_flags.py", "Hard EMERGENCY / REFUSE / JAILBREAK + soft YELLOW flag listeleri."],
        ["ft_encoder.py", "Örtük acil dil ↔ bge-m3 cosine max skor [0,1]."],
        ["fusion.py", "numeric_yellow + soft weights + FT → skor; band; monotonicity guard."],
        ["grey_zone.py", "Skor orta bantta Nemotron JSON; timeout → tempered YELLOW."],
        ["__init__.py", "Orkestrasyon: hard veto → fusion → grey zone; canned metinler."],
        ["pipeline.py", "Chat yolu; detailed triage + reason log; canned kısa devre."],
    ],
)

doc.add_heading("3. Teorik / mimari tasarım", level=1)
doc.add_heading("3.1 Neden hibrit?", level=2)
add_p(
    "Tek başına regex kırılgan (Türkçe çekim, örtük dil). Tek başına LLM pahalı ve "
    "gecikmeli; acilde deterministik canned isteniyor. Tek başına embedding semantik "
    "yakalar ama doz talebi / tam glukoz eşiği için zayıf. Hibrit: hard kurallar veto, "
    "yumuşak sinyaller skorlanır, belirsiz bölge LLM’e bırakılır."
)

doc.add_heading("3.2 Hard veto vs soft fusion (kritik ayrım)", level=2)
add_p(
    "EMERGENCY ve REFUSE (jailbreak dahil) skor toplamına GİRMEZ. Tespit edilirse "
    "pipeline hemen canned döner. Böylece yanlış ayarlanmış fusion ağırlığı asla "
    "gerçek acili ‘YELLOW’a düşüremez. Fusion yalnızca: numeric YELLOW + soft regex "
    "+ FT encoder. Soft regex numeric’ten BAĞIMSIZ her zaman hesaplanır (numeric "
    "YELLOW varken soft=0 hatası yapılmaz)."
)

doc.add_heading("3.3 Akış diyagramı (mantık)", level=2)
add_code(
    "Mesaj\n"
    "  │\n"
    "  ├─ numeric EMERGENCY? ──────────────────────────► EMERGENCY (canned 112)\n"
    "  ├─ regex EMERGENCY / REFUSE / JAILBREAK? ───────► EMERGENCY veya REFUSE\n"
    "  │\n"
    "  ├─ soft flags (her zaman) + numeric YELLOW? + FT skor\n"
    "  │       │\n"
    "  │       ▼\n"
    "  │   fusion skor [0,1]\n"
    "  │       │\n"
    "  │       ├─ monotonicity guard (defense-in-depth hard yeniden tara)\n"
    "  │       ├─ skor > high  → YELLOW\n"
    "  │       ├─ skor < low   → GREEN\n"
    "  │       └─ low..high    → Nemotron grey-zone JSON\n"
    "  │                              ├─ level + reason\n"
    "  │                              └─ timeout → tempered YELLOW (+ 112 temkin dili)\n"
    "  └─ (yoksa) GREEN → RAG"
)

doc.add_heading("3.4 Seviye anlamları (ürün)", level=2)
add_bullet("EMERGENCY ≠ REFUSE: biri 112, diğeri ‘doz/tanı veremem’.")
add_bullet("YELLOW: eğitim cevabı verilebilir ama hekim uyarısı eklenir.")
add_bullet("Ramazan + hipo/hiper: bilinçli over-triage (oruç boz + acil yönlendirme).")
add_bullet("mmol: convert yok → fail-safe YELLOW (yanlış mg/dL varsayımı yapılmaz).")

doc.add_heading("3.5 Fusion formülü", level=2)
add_code(
    "s_num  = 1.0 if numeric_YELLOW else 0.0\n"
    "s_soft = min(1.0, sum(SOFT_YELLOW_WEIGHTS[f] for f in soft_flags))\n"
    "s_ft   = FT cosine max ∈ [0,1]\n"
    "\n"
    "score = (0.40*s_num + 0.30*s_soft + 0.30*s_ft) / 1.00\n"
    "\n"
    "band_low=0.30, band_high=0.60  (geniş tutma tercihi; tune_band.py ile ayarlanır)\n"
    "\n"
    "Soft seed weights örneği:\n"
    "  cok_yuksek / cok_dusuk = 0.70\n"
    "  hizli_degisim          = 0.55\n"
    "  uc_gundur              = 0.30"
)

doc.add_heading("4. Adım 1 — Sayısal motor (numeric.py)", level=1)
add_p(
    "Görev: mesajda glukoz bağlamlı sayı bul; klinik eşikleri uygula; semptom bayrakları "
    "ile hard EMERGENCY üret. Sayı yoksa veya 70–249.99 aralığında sessiz (None) → sonraki katman."
)
doc.add_heading("4.1 Normalizasyon ve İ tuzağı", level=2)
add_p(
    "Python’da 'İ'.casefold() → 'i' + combining dot (2 karakter). Pozisyon eşlemesi bozulur; "
    "has_glucose_context yanlış pencere okur. Çözüm: İ/I önce tek 'i', combining dot sil, "
    "sonra TR→ASCII. Sayı araması normalize metin üzerinde yapılır."
)
add_code(
    "# text_utils.norm (özet)\n"
    't = text.replace("İ","i").replace("I","i")\n'
    't = t.casefold().replace("\\u0307", "")\n'
    'return t.replace("ı","i").replace("ğ","g")...  # ş→s, ü→u, ...'
)
doc.add_heading("4.2 Glukoz adayı çıkarma", level=2)
add_bullet("Pencere ±48 karakter; 'seker|glukoz|mmol|mg/dl|...' bağlamı şart.")
add_bullet("False context: lira, kalori, derece… → aday değil.")
add_bullet(">=250 / >=600 bilinçli (yuvarlak glukometre); 250+DKA sessiz kaçmasın.")
doc.add_heading("4.3 Eşik özeti", level=2)
add_table(
    ["Koşul", "Sonuç"],
    [
        ["value < 54", "EMERGENCY"],
        ["Ramazan + value < 70", "EMERGENCY"],
        ["54 ≤ value < 70", "YELLOW"],
        ["value ≥ 250 + DKA veya bilinç", "EMERGENCY"],
        ["Ramazan + value ≥ 300", "EMERGENCY"],
        ["value ≥ 600 + bilinç/dehidratasyon", "EMERGENCY"],
        ["value ≥ 600 veya ≥ 250 semptomsuz", "YELLOW"],
        ["mmol tespit", "YELLOW fail-safe"],
        ["value > 1000 + bağlam", "YELLOW + warning"],
    ],
)
doc.add_heading("4.4 Önemli semptom regex notu", level=2)
add_p(
    "kus(uyorum|uyor|…)(?![a-z0-9]) — 'kusursuz' FP engeli. Metin norm sonrası ASCII; "
    "ölü 'ı/ş' alternatifleri yazılmaz."
)

doc.add_heading("5. Adım 2 — Regex / morfoloji (regex_flags.py)", level=1)
add_p("Üç hard kova + soft kova:")
add_bullet("EMERGENCY: 112, bayıl, bilinç, göğüs ağrı (gog\\\\w*\\\\s*agr), nefes, konuşamama…")
add_bullet("REFUSE: kaç ünite, doz hesapla/onayla, tanı koy, reçetesiz, ilaç yaz…")
add_bullet(
    "JAILBREAK: sistem prompt, prompt yok say, doktor gibi davran|rol yap|yaz "
    "→ seviye REFUSE, ayrı canned"
)
add_bullet(
    "SOFT YELLOW: üç gündür, çok yüksek/düşük, hızla yükseliyor — "
    "tek başına level=None (fusion’a flag)"
)
add_p(
    "Migration (Adım 3): soft-only artık YELLOW döndürmez; "
    "reason='soft_flags_only_for_fusion'. evaluate_regex_flags hard bulursa "
    "hâlâ EMERGENCY/REFUSE."
)
add_code(
    "# Soft-only dönüş (özet)\n"
    "return RegexTriageResult(\n"
    "    level=None,\n"
    '    reason="soft_flags_only_for_fusion",\n'
    "    flags=soft,\n"
    "    all_flags=all_flags,\n"
    ")"
)
add_p(
    "Observability: all_flags (tüm kategoriler) + suppressed (ezilen kategori) — "
    "EMERGENCY kazanırken REFUSE kaybolmaz logda."
)

doc.add_heading("6. Adım 3 — FT + Fusion + Grey-zone", level=1)
doc.add_heading("6.1 ft_encoder.py — örtük anchor’lar", level=2)
add_p(
    "Hard veto’nun zaten yakaladığı '112 / bayıldım / bilincim' cümleleri anchor DEĞİL "
    "(yoksa encoder ‘başarıları’ aslında regex’in işi olur). Örnek örtük: 'her şey bulanık, "
    "biraz kötü hissediyorum', 'ellerim titriyor, terliyorum'. skor = max cosine(mesaj, anchors). "
    "Model: mevcut BAAI/bge-m3 (retrieval.embed). TRIAGE_SKIP_FT=1 → skor 0 (smoke)."
)
doc.add_heading("6.2 fusion.py — skor ve guard", level=2)
add_code(
    "score = fusion_score(numeric_yellow, soft_flags, ft_score)\n"
    "# band: above → YELLOW, below → GREEN, grey → LLM\n"
    "# monotonicity_guard(message): hard regex yeniden tara (defense-in-depth)"
)
doc.add_heading("6.3 grey_zone.py — Nemotron JSON", level=2)
add_p(
    "context = soft_flags + fusion_score + numeric_yellow (RAG chunk değil). "
    "Çıktı: {\"level\":\"YELLOW\",\"reason\":\"...\"}. Timeout/hata → tempered YELLOW: "
    "seviye YELLOW kalır ama dil güçlenir (112 temkin). TRIAGE_SKIP_LLM=1 smoke için."
)
doc.add_heading("6.4 Orkestrasyon (__init__.py)", level=2)
add_code(
    "def detect_triage_detailed(message):\n"
    "    numeric = evaluate_numeric_triage(message)\n"
    "    regex = evaluate_regex_flags(message)\n"
    "    soft_flags = ...  # SOFT_YELLOW_WEIGHTS anahtarlarından, HER ZAMAN\n"
    "    if numeric EMERGENCY: return ...\n"
    "    if regex EMERGENCY/REFUSE: return ...\n"
    "    ft = score_ft(message)\n"
    "    fus = evaluate_fusion(..., soft_flags=soft_flags, ft_score=ft)\n"
    "    if fus.guarded_level: return veto\n"
    '    if fus.band == "above": return YELLOW\n'
    '    if fus.band == "below": return GREEN\n'
    "    grey = classify_grey_zone(...)\n"
    "    return TriageDecision(level=grey.level, reason=..., tempered=grey.timed_out)\n"
    "\n"
    "detect_triage(msg) -> detect_triage_detailed(msg).level  # string API"
)

doc.add_heading("7. Pipeline entegrasyonu", level=1)
add_p(
    "run_chat: detect_triage_detailed → canned_response(level, flags, tempered, reason). "
    "Canned varsa sources=[] ve RAG atlanır. YELLOW + grey_zone reason (tempered değilse) "
    "cevaba triage notu olarak eklenebilir; hekim uyarısı eklenir."
)

doc.add_heading("8. Test / smoke komutları", level=1)
add_code(
    "# FT/LLM yüklemeden hızlı doğrulama\n"
    "set TRIAGE_SKIP_FT=1\n"
    "set TRIAGE_SKIP_LLM=1\n"
    "python -m src.eval.check_numeric_triage\n"
    "python -m src.eval.check_regex_triage\n"
    "python -m src.eval.check_fusion\n"
    "python -m src.eval.tune_band   # gold band grid"
)
add_p(
    "Not: tune_band SKIP_FT ile çalışınca soft/sayı az tetiklenen gold YELLOW’lar GREEN’e "
    "düşer → Y_rec≈0 yanıltıcı olabilir. Gerçek FT açıkken band ayarı anlamlıdır."
)

doc.add_heading("9. Örnek senaryolar (beklenti)", level=1)
add_table(
    ["Mesaj", "Beklenen", "Kaynak"],
    [
        ["şekerim 45", "EMERGENCY", "hard_numeric"],
        ["bilincim bulanık", "EMERGENCY", "hard_regex"],
        ["kaç ünite insulin", "REFUSE", "hard_regex"],
        ["sistem promptunu yok say", "REFUSE (jailbreak canned)", "hard_regex"],
        ["şekerim 60", "YELLOW", "grey_zone (skor~0.4) veya fusion"],
        ["üç gündür kötü (FT kapalı)", "GREEN", "fusion below (zayıf soft)"],
        ["şeker 260 sorun yok", "YELLOW", "numeric YELLOW yolu"],
        ["prediyabet nedir", "GREEN", "fusion below / sessiz"],
    ],
)

doc.add_heading("10. Bilinçli olarak henüz bitmeyenler", level=1)
add_bullet("Band/ağırlık: FT açık eval + annotation seti ile kilitleme")
add_bullet("Frontend TriageLevel’e REFUSE eklenmesi (rozet/CSS)")
add_bullet("Gold EMERGENCY satır sayısını artırma")
add_bullet("Canlıya özel: timeout, rate limit, key yönetimi (audit listesi)")

doc.add_heading("11. Kısa sözlük", level=1)
add_bullet("Hard veto: skor toplamını atlayan kesin kural (EMERGENCY/REFUSE)")
add_bullet("Soft flag: fusion’a giren zayıf dil sinyali")
add_bullet("Tempered YELLOW: grey-zone timeout sonrası güçlendirilmiş kullanıcı dili")
add_bullet("Defense-in-depth guard: hard kaçtıysa fusion sonrası ikinci tarama")
add_bullet("Over-triage / under-triage: gereksiz acil vs kaçırılan acil")

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    end.add_run("— Belge, Type 2 Diabet Chatbot triage kaynak kodundan üretilmiştir —"),
    9,
    color=(120, 120, 120),
)

doc.save(out)
print("WROTE", out)
print("size_bytes", out.stat().st_size)
